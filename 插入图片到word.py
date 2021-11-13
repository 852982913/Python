import os

from docx import Document

# 文档路径
filePath = "./test.docx"
print("文件保存在当下路径下：./test.docx")

print(os.listdir("."))
if filePath in os.listdir("."):
    doc = Document(filePath)
else:
    doc = Document()

# 图片路径
imgPath = input("请输入图片文件路径(文件夹,请路径确保没有中文 || 或者使用当前路径 './image'):")

try:
    for image in os.listdir(imgPath):
        print(f"正在添加图片:{os.path.join(imgPath, image)}")
        doc.add_picture(os.path.join(imgPath, image))
except Exception as e:
    print(e)
    print("路径错误")
finally:
    doc.save("test.docx")

print("添加完成")
